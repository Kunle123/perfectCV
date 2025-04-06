from typing import Any, Dict, List
from fastapi import APIRouter, Depends, HTTPException, Body
from sqlalchemy.orm import Session

from app import crud, schemas
from app.api import deps
from app.models.user import User
from app.services.resume_optimizer import optimize_resume, _extract_job_requirements

router = APIRouter()

@router.post("/optimize-resume", response_model=Dict[str, Any])
async def optimize_resume_endpoint(
    *,
    db: Session = Depends(deps.get_db),
    resume_text: str = Body(...),
    job_description: str = Body(...),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """
    Optimize a resume given its text and job description.
    """
    # Check if user has enough credits
    if current_user.credits < 1:
        raise HTTPException(
            status_code=402,
            detail="Insufficient credits. Please purchase more credits to continue."
        )
    
    # Extract job requirements
    job_requirements = await _extract_job_requirements(job_description)
    
    # Optimize resume
    optimization_result = await optimize_resume(resume_text, job_description)
    
    # Add job requirements to the result
    optimization_result["job_requirements"] = job_requirements
    
    # Create optimization record
    optimization = crud.optimization.create(
        db=db,
        obj_in=schemas.OptimizationCreate(
            resume_id=None,
            job_description_id=None,
            result=optimization_result,
            status="completed"
        )
    )
    
    # Deduct credits
    current_user.credits -= 1
    db.commit()
    
    return optimization_result

@router.get("/", response_model=List[schemas.Optimization])
def read_optimizations(
    db: Session = Depends(deps.get_db),
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """
    Retrieve optimizations.
    """
    optimizations = crud.optimization.get_multi_by_user(
        db=db, user_id=current_user.id, skip=skip, limit=limit
    )
    return optimizations

@router.get("/{optimization_id}", response_model=schemas.Optimization)
def read_optimization(
    *,
    db: Session = Depends(deps.get_db),
    optimization_id: int,
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """
    Get optimization by ID.
    """
    optimization = crud.optimization.get(db=db, id=optimization_id)
    if not optimization:
        raise HTTPException(status_code=404, detail="Optimization not found")
    if optimization.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    return optimization

@router.delete("/{optimization_id}", response_model=schemas.Optimization)
def delete_optimization(
    *,
    db: Session = Depends(deps.get_db),
    optimization_id: int,
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """
    Delete optimization.
    """
    optimization = crud.optimization.get(db=db, id=optimization_id)
    if not optimization:
        raise HTTPException(status_code=404, detail="Optimization not found")
    if optimization.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    optimization = crud.optimization.remove(db=db, id=optimization_id)
    return optimization

@router.post("/export/{optimization_id}")
async def export_optimization(
    *,
    db: Session = Depends(deps.get_db),
    optimization_id: int,
    format: str = Body("pdf"),
    current_user: User = Depends(deps.get_current_active_user),
) -> Any:
    """
    Export optimized resume in specified format.
    """
    optimization = crud.optimization.get(db=db, id=optimization_id)
    if not optimization:
        raise HTTPException(status_code=404, detail="Optimization not found")
    if optimization.resume.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    
    # Generate file based on requested format
    from fastapi.responses import FileResponse
    import tempfile
    import os
    from jinja2 import Template
    import pdfkit
    from docx import Document
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}") as temp_file:
        temp_path = temp_file.name
    
    # Get optimization result data
    result_data = optimization.result
    
    if format.lower() == "pdf":
        # Create HTML template
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Optimized Resume</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { color: #2c3e50; }
                h2 { color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }
                .section { margin-bottom: 20px; }
                .suggestion { margin-bottom: 10px; padding: 5px; background-color: #f8f9fa; }
            </style>
        </head>
        <body>
            <h1>Optimized Resume</h1>
            <div class="section">
                <h2>Optimization Suggestions</h2>
                {% for section in suggestions %}
                <h3>{{ section.section }}</h3>
                {% for suggestion in section.suggestions %}
                <div class="suggestion">{{ suggestion }}</div>
                {% endfor %}
                {% endfor %}
            </div>
        </body>
        </html>
        """
        
        # Render template with data
        template = Template(html_template)
        html_content = template.render(suggestions=result_data.get("suggestions", []))
        
        # Write HTML to temporary file
        html_path = temp_path.replace(".pdf", ".html")
        with open(html_path, "w") as html_file:
            html_file.write(html_content)
        
        # Convert HTML to PDF
        pdfkit.from_file(html_path, temp_path)
        
        # Clean up HTML file
        os.unlink(html_path)
        
    elif format.lower() == "docx":
        # Create Word document
        doc = Document()
        doc.add_heading('Optimized Resume', 0)
        
        # Add optimization suggestions
        doc.add_heading('Optimization Suggestions', level=1)
        for section in result_data.get("suggestions", []):
            doc.add_heading(section.get("section", ""), level=2)
            for suggestion in section.get("suggestions", []):
                doc.add_paragraph(suggestion)
        
        # Save document
        doc.save(temp_path)
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
    
    # Return the file as a response
    return FileResponse(
        path=temp_path,
        filename=f"optimized_resume_{optimization_id}.{format}",
        media_type=f"application/{format}"
    )